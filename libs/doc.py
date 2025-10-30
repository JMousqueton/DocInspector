import os
from docx import Document

def is_docx_file(filename):
    # True if file is .docx and starts with PK (zip signature)
    return filename.lower().endswith('.docx') and _has_zip_sig(filename)

def _has_zip_sig(filename):
    try:
        with open(filename, "rb") as f:
            return f.read(2) == b"PK"
    except Exception:
        return False

def extract_docx_metadata(docx_file):
    doc = Document(docx_file)
    core = doc.core_properties
    meta = [
        ["title", core.title or ""],
        ["subject", core.subject or ""],
        ["creator", core.author or ""],
        ["keywords", core.keywords or ""],
        ["description", core.comments or ""],
        ["last_modified_by", core.last_modified_by or ""],
        ["revision", core.revision or ""],
        ["created", str(core.created) if core.created else ""],
        ["modified", str(core.modified) if core.modified else ""],
        ["category", core.category or ""],
        ["content_status", core.content_status or ""],
        ["identifier", core.identifier or ""],
        ["language", core.language or ""],
        ["version", core.version or ""]
    ]
    # Custom properties (if any)
    if hasattr(doc, 'custom_properties'):
        for key in doc.custom_properties:
            meta.append([f"custom_{key}", str(doc.custom_properties[key])])
    # Optionally extract the template property (from app.xml)
    template = get_docx_template_name(docx_file)
    if template:
        meta.append(["template", template])
    return meta

def extract_docx_links(docx_file):
    from docx import Document
    doc = Document(docx_file)
    links = set()
    # Hyperlinks in paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if hasattr(run, "hyperlink") and run.hyperlink and run.hyperlink.target:
                links.add(run.hyperlink.target)
    # Tables (optional: you may want to extract from cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if hasattr(run, "hyperlink") and run.hyperlink and run.hyperlink.target:
                            links.add(run.hyperlink.target)
    return sorted(links)

def extract_docx_images(docx_file):
    from zipfile import ZipFile
    images = []
    with ZipFile(docx_file) as zf:
        for name in zf.namelist():
            if name.startswith('word/media/'):
                images.append(name)
    return images

def extract_docx_comments(docx_file):
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET
    comments = []
    try:
        with ZipFile(docx_file) as docx_zip:
            for name in docx_zip.namelist():
                if name == 'word/comments.xml':
                    with docx_zip.open(name) as f:
                        root = ET.parse(f).getroot()
                        for comment in root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'):
                            author = comment.attrib.get('author', '')
                            date = comment.attrib.get('date', '')
                            text = ''.join(child.text or '' for child in comment.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
                            comments.append({'author': author, 'date': date, 'text': text})
    except Exception:
        pass
    return comments

def has_vba_macros(docx_file):
    """
    Returns True if word/vbaProject.bin is present, indicating VBA macros.
    """
    from zipfile import ZipFile
    with ZipFile(docx_file) as docx_zip:
        return "word/vbaProject.bin" in docx_zip.namelist()

def extract_custom_xml_parts(docx_file):
    """
    Returns a list of custom XML part filenames and (optionally) their contents.
    """
    from zipfile import ZipFile
    xml_parts = []
    with ZipFile(docx_file) as docx_zip:
        for name in docx_zip.namelist():
            if name.startswith("customXml/") and name.endswith(".xml"):
                try:
                    with docx_zip.open(name) as f:
                        content = f.read().decode("utf-8", errors="replace")
                    xml_parts.append({"filename": name, "content": content})
                except Exception:
                    xml_parts.append({"filename": name, "content": "(unreadable)"})
    return xml_parts

def get_docx_template_name(docx_file):
    # Try to extract template property from docProps/app.xml
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET
    try:
        with ZipFile(docx_file) as docx_zip:
            with docx_zip.open("docProps/app.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                ns = {'ap': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                template_elem = root.find('ap:Template', ns)
                if template_elem is not None:
                    return template_elem.text
    except Exception:
        pass
    return ""

# Dummy for .doc files (not supported)
def get_doc_basic_info(filename):
    return {
        "file_size_bytes": None,
        "file_size_human": "",
        "num_paragraphs": 0,
        "num_tables": 0,
        "meta": [["error", ".doc (legacy Word) not supported in this script"]],
        "links": [],
        "images": [],
        "comments": [],
        "has_vba_macros": False,
        "custom_xml_parts": [],
    }

def get_docx_num_pages(docx_file):
    """
    Attempts to read the number of pages from docProps/app.xml
    Returns an integer or None if not available.
    """
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET
    try:
        with ZipFile(docx_file) as docx_zip:
            with docx_zip.open("docProps/app.xml") as f:
                root = ET.parse(f).getroot()
                ns = {'ap': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                pages_elem = root.find('ap:Pages', ns)
                if pages_elem is not None:
                    return int(pages_elem.text)
    except Exception:
        pass
    return None

def has_revision_marks(docx_file):
    """
    Returns True if <w:ins>, <w:del>, <w:moveFrom>, or <w:moveTo> elements are present.
    """
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET
    try:
        with ZipFile(docx_file) as docx_zip:
            if "word/document.xml" in docx_zip.namelist():
                xml = docx_zip.read("word/document.xml")
                root = ET.fromstring(xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for tag in ("ins", "del", "moveFrom", "moveTo"):
                    if root.find('.//w:' + tag, ns) is not None:
                        return True
    except Exception:
        pass
    return False

def get_docx_basic_info(docx_file):
    from libs.shared import human_readable_size
    file_size = os.path.getsize(docx_file)
    doc = Document(docx_file)
    meta = extract_docx_metadata(docx_file)
    num_paragraphs = len(doc.paragraphs)
    num_tables = len(doc.tables)
    num_pages = get_docx_num_pages(docx_file)
    links = extract_docx_links(docx_file)
    images = extract_docx_images(docx_file)
    comments = extract_docx_comments(docx_file)
    has_macros = has_vba_macros(docx_file)
    custom_xml_parts = extract_custom_xml_parts(docx_file)
    revision_marks = has_revision_marks(docx_file)
    return {
        "file_size_bytes": file_size,
        "file_size_human": human_readable_size(file_size),
        "num_pages": num_pages,
        "num_paragraphs": num_paragraphs,
        "num_tables": num_tables,
        "meta": meta,
        "links": links,
        "images": images,
        "comments": comments,
        "has_vba_macros": has_macros,
        "custom_xml_parts": custom_xml_parts,
        "has_revision_marks": revision_marks,
    }
